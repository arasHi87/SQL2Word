import sys
import mysql.connector

from logger import logger
from docx import Document
from pprint import pprint
from collections import OrderedDict

class ConvertSQL2Word():
    def __init__(self, host, user, passwd,
                 db, ssl_ca, ssl_cert, ssl_key):
        self.db = mysql.connector.connect(
            host = host,
            user = user,
            passwd = passwd,
            db = db,
            ssl_ca = ssl_ca,
            ssl_cert = ssl_cert,
            ssl_key = ssl_key)
        self.cursor = self.db.cursor()
    
    def getTables(self, t_name):
        sql = """
        SELECT TABLE_NAME
        FROM INFORMATION_SCHEMA.TABLES
        WHERE TABLE_TYPE = 'BASE TABLE' AND TABLE_SCHEMA = '{}'""".format(t_name)
        tables = OrderedDict()

        self.cursor.execute(sql)
        
        for x in self.cursor.fetchall():
            logger.info('get table {}'.format(x[0]))

            tables[x[0]] = list()

        return tables

    def getColumns(self, table, db):
        logger.info('get foreign keys with {}'.format(table))

        sql = """
        SELECT COLUMN_NAME
        FROM INFORMATION_SCHEMA.KEY_COLUMN_USAGE
        WHERE REFERENCED_TABLE_SCHEMA = '{}'
        AND TABLE_NAME = '{}'
        AND CONSTRAINT_NAME LIKE '%foreign%'""".format(db, table)
        foreign_keys = dict()
        
        self.cursor.execute(sql)

        for x in self.cursor.fetchall():
            foreign_keys[x[0]] = 1
        
        logger.info('Starting insert column info into {}'.format(table))
    
        sql = 'SHOW FULL COLUMNS FROM `{}`'.format(table)
        columns = list()
        
        self.cursor.execute(sql)

        result = self.cursor.fetchall()

        for x in result:
            column = dict()
            column[x[0]] = dict()
            column[x[0]]['field'] = x[0]
            column[x[0]]['type'] = x[1]
            column[x[0]]['null'] = x[3][0]
            column[x[0]]['default'] = x[5]
            column[x[0]]['pk'] = 'N'
            column[x[0]]['fk'] = 'N'
            column[x[0]]['ik'] = 'N'
            column[x[0]]['mk'] = 'N'
            column[x[0]]['comment'] = x[8]

            if x[4] == 'PRI': column[x[0]]['pk'] = 'Y'
            if x[4] == 'MUL': column[x[0]]['mk'] = 'Y'
            if x[0] in foreign_keys.keys(): column[x[0]]['fk'] = 'Y'
            if column[x[0]]['null'] == 'Y': column[x[0]]['default'] = 'NULL'
            elif not column[x[0]]['default']: column[x[0]]['default'] = ''
        
            columns.append(column)

        return columns
    
    def convert(self, data):
        logger.info('Starting to convert sql to word')

        document = Document()
        styles = document.styles

        for key, columns in data.items():
            document.add_heading('Table: {} ()'.format(key), level=1)
            document.add_paragraph()
            
            table = document.add_table(rows=1, cols=8)

            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Type'
            hdr_cells[2].text = 'Null'
            hdr_cells[3].text = 'Default'
            hdr_cells[4].text = 'PK'
            hdr_cells[5].text = 'IK'
            hdr_cells[6].text = 'FK'
            hdr_cells[7].text = 'Comment' 

            for column in columns:
                for idx, value in column.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = value['field']
                    row_cells[1].text = value['type']
                    row_cells[2].text = value['null']
                    row_cells[3].text = value['default']
                    row_cells[4].text = value['pk']
                    row_cells[5].text = value['ik']
                    row_cells[6].text = value['fk']
                    row_cells[7].text = value['comment'] 

            document.add_page_break()

        document.save('meow.docx')
        logger.info('Success convert')

if __name__ == '__main__':
    args = sys.argv
    host = args[1]
    user = args[2]
    passwd = args[3]
    db = args[4]
    ssl_ca = args[5]
    ssl_cert = args[6]
    ssl_key = args[7]
    s2w = ConvertSQL2Word(
        host = host,
        user = user,
        passwd =  passwd,
        db = db,
        ssl_ca = ssl_ca,
        ssl_cert = ssl_cert,
        ssl_key = ssl_key
    )
    tables = s2w.getTables(db) 

    for key, value in tables.items():
        tables[key] = s2w.getColumns(key, db) 

    s2w.convert(tables)
